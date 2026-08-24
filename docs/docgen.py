"""Shared helpers for building the DSR Portal Word documents.

Both deliverables are generated from code so they can be regenerated after a
release without hand-editing: screenshots refresh, headings stay consistent,
and the styling is defined once here.
"""
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).resolve().parent.parent
SHOTS = ROOT / "tools" / "docshots"

INK = RGBColor(0x0A, 0x0A, 0x0A)
MUTED = RGBColor(0x55, 0x55, 0x55)
# Brand yellow is a fill, not an ink: on white paper it is 1.2:1, so headings
# take the dark gold instead, matching --t-brand-ink in the console.
BRAND = RGBColor(0x6F, 0x59, 0x00)
FAINT = RGBColor(0x8A, 0x8A, 0x85)


def new_document(title: str, subtitle: str, audience: str, version: str) -> Document:
    """A document with the house styles and a title page already applied."""
    doc = Document()

    base = doc.styles["Normal"]
    base.font.name = "Calibri"
    base.font.size = Pt(10.5)
    base.font.color.rgb = INK
    base.paragraph_format.space_after = Pt(7)
    base.paragraph_format.line_spacing = 1.18

    for name, size, color, before, after in [
        ("Heading 1", 20, INK, 22, 8),
        ("Heading 2", 14.5, INK, 18, 6),
        ("Heading 3", 11.5, BRAND, 12, 4),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True
        st.paragraph_format.space_before = Pt(before)
        st.paragraph_format.space_after = Pt(after)
        st.paragraph_format.keep_with_next = True

    code = doc.styles.add_style("CodeBlock", 1)  # WD_STYLE_TYPE.PARAGRAPH
    code.font.name = "Consolas"
    code.font.size = Pt(8.5)
    code.paragraph_format.space_before = Pt(4)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.left_indent = Inches(0.16)

    cap = doc.styles.add_style("Caption2", 1)
    cap.font.name = "Calibri"
    cap.font.size = Pt(8.5)
    cap.font.italic = True
    cap.font.color.rgb = FAINT
    cap.paragraph_format.space_before = Pt(2)
    cap.paragraph_format.space_after = Pt(12)

    for section in doc.sections:
        section.top_margin = Inches(0.85)
        section.bottom_margin = Inches(0.85)
        section.left_margin = Inches(0.9)
        section.right_margin = Inches(0.9)

    # ------------------------------------------------------------ title page
    for _ in range(5):
        doc.add_paragraph()

    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = t.add_run(title)
    run.font.size = Pt(30)
    run.font.bold = True
    run.font.color.rgb = INK

    s = doc.add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = s.add_run(subtitle)
    run.font.size = Pt(13)
    run.font.color.rgb = MUTED

    doc.add_paragraph()
    a = doc.add_paragraph()
    a.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = a.add_run(audience)
    run.font.size = Pt(10)
    run.font.color.rgb = BRAND
    run.font.bold = True

    v = doc.add_paragraph()
    v.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = v.add_run(version)
    run.font.size = Pt(9)
    run.font.color.rgb = FAINT

    doc.add_page_break()
    return doc


def add_toc(doc: Document) -> None:
    """Insert a Word field for a table of contents.

    Word populates it on open; some readers ask the user to update fields
    first, which is why the instruction line sits underneath.
    """
    doc.add_heading("Contents", level=1)
    p = doc.add_paragraph()
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), r'TOC \o "1-3" \h \z \u')
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = "Right-click here and choose Update Field to build the contents."
    run.append(text)
    fld.append(run)
    p._p.append(fld)

    note = doc.add_paragraph(
        "If the list above is empty, click it once and press F9 to populate it."
    )
    note.style = doc.styles["Caption2"]
    doc.add_page_break()


def h1(doc, text):
    doc.add_heading(text, level=1)


def h2(doc, text):
    doc.add_heading(text, level=2)


def h3(doc, text):
    doc.add_heading(text, level=3)


def para(doc, text, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    return p


def bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        _rich(p, it)


def numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        _rich(p, it)


def _rich(paragraph, text):
    """Render **bold** and `code` spans inside a paragraph."""
    import re

    for part in re.split(r"(\*\*[^*]+\*\*|`[^`]+`)", text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = paragraph.add_run(part[2:-2])
            run.bold = True
        elif part.startswith("`") and part.endswith("`"):
            run = paragraph.add_run(part[1:-1])
            run.font.name = "Consolas"
            run.font.size = Pt(9)
            run.font.color.rgb = BRAND
        else:
            paragraph.add_run(part)


def rich_para(doc, text):
    p = doc.add_paragraph()
    _rich(p, text)
    return p


def code(doc, text):
    for line in text.strip("\n").split("\n"):
        p = doc.add_paragraph(line or " ", style="CodeBlock")
    return p


def callout(doc, kind, text):
    """A shaded single-cell table used for notes and warnings."""
    colors = {"note": "EEF0FB", "warn": "FDF3E7", "danger": "FCEDEC", "ok": "EDF7F0"}
    labels = {"note": "Note", "warn": "Important", "danger": "Warning", "ok": "Tip"}
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.rows[0].cells[0]
    shade = OxmlElement("w:shd")
    shade.set(qn("w:fill"), colors[kind])
    cell._tc.get_or_add_tcPr().append(shade)
    p = cell.paragraphs[0]
    run = p.add_run(f"{labels[kind]}  ")
    run.bold = True
    run.font.size = Pt(9.5)
    _rich(p, text)
    for r in p.runs[1:]:
        r.font.size = Pt(9.5)
    doc.add_paragraph()


def table(doc, headers, rows, widths=None):
    t = doc.add_table(rows=1, cols=len(headers))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, head in enumerate(headers):
        hdr[i].text = ""
        run = hdr[i].paragraphs[0].add_run(head)
        run.bold = True
        run.font.size = Pt(9)
    for row in rows:
        cells = t.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = ""
            p = cells[i].paragraphs[0]
            _rich(p, str(value))
            for r in p.runs:
                r.font.size = Pt(9)
    if widths:
        for row in t.rows:
            for i, w in enumerate(widths):
                row.cells[i].width = Inches(w)
    doc.add_paragraph()
    return t


def figure(doc, name, caption, width=6.4):
    """Insert a screenshot with a caption; skip quietly if it is missing."""
    path = SHOTS / f"{name}.png"
    if not path.exists():
        return
    doc.add_picture(str(path), width=Inches(width))
    doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap = doc.add_paragraph(caption)
    cap.style = doc.styles["Caption2"]
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def page_break(doc):
    doc.add_page_break()
